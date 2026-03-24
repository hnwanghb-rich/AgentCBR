from docx import Document
from docx.shared import Inches
from pathlib import Path
from datetime import datetime

class OutputGenerator:
    def __init__(self, output_dir, logger):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.logger = logger

    def generate_document(self, title, content, attachments, index):
        try:
            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(f"记录编号: {index}")
            doc.add_paragraph(f"处理时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_heading("内容", level=1)
            doc.add_paragraph(content)

            if attachments:
                doc.add_heading("附件内容", level=1)
                for att in attachments:
                    if att['type'] == 'text':
                        doc.add_paragraph(att['content'])
                    elif att['type'] == 'image' and att.get('image'):
                        img_path = f"temp_img_{index}.png"
                        att['image'].save(img_path)
                        doc.add_picture(img_path, width=Inches(5))
                        doc.add_paragraph(f"图片识别文字: {att['content']}")

            filename = self.output_dir / f"record_{index}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)
            self.logger.info(f"文档已生成: {filename}")
        except Exception as e:
            self.logger.error(f"生成文档失败: {e}")

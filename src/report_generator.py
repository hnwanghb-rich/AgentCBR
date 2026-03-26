"""碳盘查报告生成模块 - Word格式"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List
from pathlib import Path
import time

class ReportGenerator:
    """碳盘查报告生成器"""

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_report(self, data: Dict, issues: List[Dict], filename: str = None) -> str:
        """生成碳盘查报告"""
        if not filename:
            filename = f"碳盘查报告_{time.strftime('%Y%m%d_%H%M%S')}.docx"

        filepath = self.output_dir / filename
        doc = Document()

        # 标题
        title = doc.add_heading('碳盘查审核报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息
        doc.add_heading('一、项目基本信息', level=1)
        self._add_basic_info(doc, data)

        # 审核结果汇总
        doc.add_heading('二、审核结果汇总', level=1)
        self._add_summary(doc, issues)

        # 问题清单
        doc.add_heading('三、问题清单', level=1)
        self._add_issue_table(doc, issues)

        doc.save(filepath)
        return str(filepath)

    def _add_basic_info(self, doc, data: Dict):
        """添加基本信息"""
        info = [
            f"公司名称：{data.get('company', '')}",
            f"项目名称：{data.get('project', '')}",
            f"审核日期：{time.strftime('%Y年%m月%d日')}",
            f"审核人员：系统自动审核"
        ]
        for line in info:
            doc.add_paragraph(line)

    def _add_summary(self, doc, issues: List[Dict]):
        """添加审核结果汇总"""
        total = len(issues)
        high = sum(1 for i in issues if i.get('priority') == '高')
        medium = sum(1 for i in issues if i.get('priority') == '中')
        low = sum(1 for i in issues if i.get('priority') == '低')

        doc.add_paragraph(f"共发现 {total} 个问题，其中：")
        doc.add_paragraph(f"  高优先级：{high} 个")
        doc.add_paragraph(f"  中优先级：{medium} 个")
        doc.add_paragraph(f"  低优先级：{low} 个")

    def _add_issue_table(self, doc, issues: List[Dict]):
        """添加问题清单表格"""
        if not issues:
            doc.add_paragraph("未发现问题")
            return

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        headers = ['序号', '项目名称', '问题描述', '修改意见', '优先级']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for idx, issue in enumerate(issues, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = issue.get('project', '')
            row.cells[2].text = issue.get('description', '')
            row.cells[3].text = issue.get('suggestion', '')
            row.cells[4].text = issue.get('priority', '')

